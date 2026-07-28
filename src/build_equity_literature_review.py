"""
Builds output/Mobius_Wealth_Equity_Income_Literature_Review.docx - internship Week 5, Task 11
("Review literature on using equities for retirement income and identify why equities may be
suitable for this purpose"). Styled to match output/Mobius_Wealth_Methodology_and_Assumptions.docx
(same title/heading colours and sizes) so the two read as one consistent set of project documents.

Run: `python build_equity_literature_review.py`
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent.parent / "output" / "Mobius_Wealth_Equity_Income_Literature_Review.docx"

NAVY = RGBColor(0x1A, 0x2B, 0x3C)
GREEN = RGBColor(0x1B, 0xAF, 0x7A)
GREY = RGBColor(0x5A, 0x5A, 0x5A)

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def title_block():
    p = doc.add_paragraph()
    r = p.add_run("Mobius Wealth")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    r = p.add_run("Equities for Retirement Income")
    r.font.size = Pt(20)
    r.font.color.rgb = GREEN

    p = doc.add_paragraph()
    r = p.add_run("Literature Review")
    r.font.size = Pt(16)
    r.font.color.rgb = GREY
    r2 = p.add_run("\nInternship Week 5, Task 11 - grounds the design of the Weeks 5-8 individual-"
                   "share/basket decumulation framework (src/equity_income.py) in existing research, "
                   "rather than testing shares against probability-of-ruin with no underlying rationale.")
    r2.font.size = Pt(11)
    doc.add_paragraph()


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def body(text):
    doc.add_paragraph(text)


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


title_block()

h1("Contents")
for i, t in enumerate([
    "Why consider equities for retirement income at all",
    "Withdrawal-rate literature: the 4% rule and the Trinity study",
    "Sequence-of-returns risk: why the first decade dominates the outcome",
    "Equity glide paths: should allocation rise or fall through retirement?",
    "Dividend income vs total return: two competing philosophies",
    "Concentration risk: a UK-specific caution",
    "What this means for this project's methodology",
    "References",
], start=1):
    bullet(f"{i}. {t}")
doc.add_paragraph()

h1("1. Why consider equities for retirement income at all")
body(
    "A retirement pot commonly needs to fund 20-30+ years of spending. Cash and short bonds are "
    "capital-stable but historically struggle to outpace inflation over that length of horizon, "
    "which erodes real spending power exactly when a plan is judged over decades rather than years. "
    "Jeremy Siegel's long-run analysis of two centuries of US market data (Stocks for the Long Run) "
    "found real (inflation-adjusted) equity returns of roughly 6.5-7% pa, materially ahead of bonds "
    "and cash over the same long sweep - the 'equity risk premium' investors are compensated with "
    "for bearing higher short-term volatility."
)
body(
    "Two caveats worth carrying into this project rather than treating the equity risk premium as "
    "guaranteed: recent academic work has questioned how constant that premium really is once "
    "different data sources and sub-periods are examined, and stocks are not a perfect short-to-"
    "medium-term inflation hedge even though they have tended to beat inflation over long horizons. "
    "Both points support this project's existing approach of testing outcomes via a full historical "
    "bootstrap (and the optional forward-looking CMA blend already in the main simulator) rather than "
    "assuming a single fixed expected equity return."
)
body(
    "This is also why every portfolio already modelled in the main Mobius Wealth simulator (Original, "
    "Alternative, Four Seasons, Better) carries meaningful equity weight throughout decumulation "
    "rather than de-risking entirely into bonds/cash - the literature above is the underlying reason, "
    "not just a modelling convention."
)

h1("2. Withdrawal-rate literature: the 4% rule and the Trinity study")
body(
    "William Bengen's 1994 paper, 'Determining Withdrawal Rates Using Historical Data' (Journal of "
    "Financial Planning), is the origin of the '4% rule'. Using rolling 30-year US retirement periods "
    "since 1926 and a 60% stock / 40% intermediate-Treasury mix rebalanced annually, Bengen found a "
    "maximum sustainable first-year withdrawal rate of about 4.15% (rounded down to 4% in later "
    "commentary) - notably, this used a majority-equity portfolio, not a conservative one."
)
body(
    "The 1998 Trinity study (Cooley, Hubbard & Walz, Journal of the American Association of "
    "Individual Investors) followed the same historical-simulation methodology but reframed the "
    "question around 'portfolio success rate' (percentage of historical 30-year periods a portfolio "
    "survived) across five stock/bond mixes from 100% equity to 100% bonds. At a 4% withdrawal rate, "
    "50% and 75% equity allocations achieved roughly 95% and 98% success rates respectively over "
    "30-year periods - i.e. adding equity exposure improved historical sustainability at that "
    "withdrawal rate, it did not just add risk for no benefit."
)
body(
    "Both studies are US-market, US-inflation studies. The Mobius Wealth simulator already runs the "
    "equivalent test on UK/global asset-class data; extending the same logic down to individual UK "
    "shares is what Weeks 5-8 (and equity_income.py) do next."
)

h1("3. Sequence-of-returns risk: why the first decade dominates the outcome")
body(
    "Two retirees can experience identical average returns over 30 years and end up with very "
    "different outcomes purely because of the ORDER those returns arrive in. A market fall early in "
    "retirement forces more units/shares to be sold to fund the same cash withdrawal, permanently "
    "reducing the capital left to participate in the eventual recovery - a loss that a market fall of "
    "the same size late in retirement does not inflict in the same way, since by then either less "
    "capital remains at risk or the plan is closer to its end regardless."
)
body(
    "Wade Pfau's research estimates that roughly 77% of a retirement portfolio's final outcome can be "
    "explained by the returns of just the first decade of retirement. Industry analysis built on the "
    "same historical simulations commonly attributes the large majority of retirement-plan failures to "
    "negative returns occurring in the first five years specifically."
)
body(
    "This is precisely why probability of ruin - not average or median return - is the headline metric "
    "throughout the main Mobius Wealth simulator, and why the stress-scenario tooling already built "
    "into the app (dot-com crash, Global Financial Crisis, 2022 inflation shock starting points) tests "
    "specific bad-sequence starting points explicitly, rather than only reporting an averaged, "
    "sequence-blind outcome."
)

h1("4. Equity glide paths: should allocation rise or fall through retirement?")
body(
    "Conventional advice (e.g. '110 minus age' style rules) reduces equity exposure as retirement "
    "approaches and progresses, on the assumption that a shorter remaining horizon should mean less "
    "risk. Pfau & Kitces' 2013-2014 research ('Reducing Retirement Risk with a Rising Equity Glide-"
    "Path', and the 2015 follow-up 'Retirement Risk, Rising Equity Glidepaths, and Valuation-Based "
    "Asset Allocation') tested this against US historical data and found the opposite can hold: a "
    "RISING equity glide path through retirement - starting more conservative (roughly 20-40% equity) "
    "and finishing more aggressive (roughly 40-80%) - reduced both the probability and the severity of "
    "plan failure versus a flat allocation or a traditional declining glide path in their tests."
)
body(
    "The logic is sequence-of-returns risk again, viewed from the other direction: a rising glide path "
    "keeps the retiree LESS exposed to loss during the most vulnerable early years (Section 3), while "
    "still allowing meaningful equity exposure to be in place once markets eventually recover and the "
    "highest-risk window has passed."
)
body(
    "The main simulator's existing Glide Path tab already tests exactly this question (de-risking vs "
    "up-risking equity weight over the horizon) at the portfolio level. Task 16 (explore weighting/"
    "rebalancing approaches) extends the same underlying question to individual-share baskets via "
    "equity_income.py's three rebalancing conventions - constant-mix (monthly), annual, and "
    "buy-and-hold - which is a related but distinct question (how a FIXED target mix is maintained) "
    "from whether the target mix itself should change over time."
)

h1("5. Dividend income vs total return: two competing philosophies")
body(
    "Two schools of thought exist for funding retirement spending from equities. The 'natural yield' "
    "approach spends only the dividend income a portfolio produces and leaves the capital untouched. "
    "The 'total return' approach holds a portfolio (often more diversified and lower-yield) for its "
    "combined income-plus-growth return, and periodically sells holdings to fund spending regardless "
    "of whether that return arrived as a dividend or as capital appreciation."
)
body(
    "The practitioner case for total return: restricting a portfolio to high-yield shares concentrates "
    "it into value-tilted, income-paying sectors and away from the broader market (see Section 6), and "
    "a total-return approach that chooses what to sell and when can be more tax-efficient than passively "
    "accepting whatever dividend a company happens to declare. The case for natural yield is largely "
    "behavioural - 'never touch the capital' is simple to explain and to stick to, and the income "
    "continues regardless of short-term price moves so long as the dividend itself is not cut. Reviewed "
    "sources broadly agree neither approach is universally superior; a hybrid tilt towards quality, "
    "growing dividend payers (rather than simply the highest current yield) is a common middle ground."
)
body(
    "This is a design choice worth stating explicitly rather than leaving implicit: equity_income.py "
    "evaluates every share and basket via the SAME Monte Carlo pot/withdrawal mechanics as the rest of "
    "this project (run_simulation, downside_stats) - spending is funded from total portfolio value, "
    "irrespective of whether the underlying return came from dividends or price growth. This project is "
    "therefore taking the total-return view of Section 5, not a 'spend the natural yield only' view."
)

h1("6. Concentration risk: a UK-specific caution")
body(
    "The UK market is unusually concentrated for dividend income: historically around a third of all "
    "FTSE All-Share dividends have come from just five companies, several of which (HSBC, the former "
    "Royal Dutch Shell, Lloyds) cut or cancelled their dividends in 2020. Headline UK plc dividends fell "
    "roughly 44% that year to their lowest level since 2011, and the average IA UK Equity Income fund "
    "still saw payouts fall around 29% despite active management - severe enough that the Investment "
    "Association temporarily suspended its normal 90%-yield sector-eligibility rule."
)
body(
    "This sits alongside a more general point: individual constituents of a broad index run "
    "substantially higher volatility than the index itself (one commonly cited comparison: US large-"
    "cap constituents of the Russell 1000 have averaged roughly 37% annualised volatility since 2014, "
    "versus about 15% for the index as a whole) - the same logic applies to individual UK shares tested "
    "one at a time."
)
body(
    "This is precisely why equity_income.py does not stop at ranking single shares by probability of "
    "ruin (task 13). It goes on to systematically search combinations (task 14's find_best_baskets, "
    "which brute-forces every equal-weight combination rather than hand-picking a plausible-looking "
    "pair) and to check pairwise correlation (share_correlation_matrix) before treating any basket as "
    "genuinely diversified - a single share with an attractive backtested probability of ruin can still "
    "carry concentration and single-company risk that a properly diversified basket does not."
)

h1("7. What this means for this project's methodology")
body(
    "Each design choice already built into equity_income.py traces back to a specific point above, "
    "rather than being an arbitrary modelling decision:"
)
bullet("Probability of ruin as the headline objective, not average return - Section 3 (sequence risk).")
bullet(
    "Testing individual shares AND systematic basket combinations, backed by a correlation matrix - "
    "Section 6 (concentration risk)."
)
bullet(
    "Three rebalancing conventions (constant-mix/monthly, annual, buy-and-hold) - Section 4 "
    "(allocation-path research), extended from the portfolio-level Glide Path question to individual "
    "share baskets."
)
bullet(
    "Total-return spending mechanics (spend from total pot value, not dividend income only) - Section 5, "
    "an explicit, stated choice rather than an oversight."
)
body(
    "The one gap this literature review does not close: every result equity_income.py currently "
    "produces runs on SYNTHETIC placeholder share data (generate_placeholder_equity_data.py), built to "
    "prove the mechanics work, not to represent real UK companies. Task 12 - pulling real UK equity "
    "data from Bloomberg - is what turns this from a validated framework into a real finding, and is "
    "covered in the companion document, 'UK Equity Data Extraction Spec (Task 12)'."
)

h1("8. References")
refs = [
    "Bengen, W.P. (1994). 'Determining Withdrawal Rates Using Historical Data.' Journal of Financial Planning.",
    "Cooley, P.L., Hubbard, C.M. & Walz, D.T. (1998). 'Retirement Spending: Choosing a Sustainable "
    "Withdrawal Rate.' Journal of the American Association of Individual Investors (the 'Trinity study').",
    "Pfau, W.D. & Kitces, M.E. (2013/2014). 'Reducing Retirement Risk with a Rising Equity Glide-Path.'",
    "Kitces, M.E. & Pfau, W.D. (2015). 'Retirement Risk, Rising Equity Glidepaths, and Valuation-Based "
    "Asset Allocation.' Journal of Financial Planning.",
    "Siegel, J.J. Stocks for the Long Run.",
    "Trustnet (2021). 'The funds that dodged the IA UK Equity Income sector's 29.3% dividend cut.' "
    "https://www.trustnet.com/News/7466862",
    "Portfolio Adviser. 'Is the equity income model broken?' https://portfolio-adviser.com/"
    "is-the-old-equity-income-model-broken/",
    "Charles Schwab. 'Using a Total-Return Approach to Retirement Income.' "
    "https://www.schwab.com/learn/story/how-to-use-total-return-approach-retirement-income",
    "Retirement Researcher. 'Why Sequence of Return Risk Matters for Your Retirement Income.' "
    "https://retirementresearcher.com/why-sequence-of-return-risk-matters-for-your-retirement-income/",
]
for r in refs:
    bullet(r)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run(
    "Note on sourcing: academic papers above (Bengen; Cooley, Hubbard & Walz; Pfau & Kitces) are "
    "well-established, widely-cited retirement-planning research summarised here from secondary "
    "coverage rather than the original journal PDFs; industry/press sources (Trustnet, Portfolio "
    "Adviser, Schwab, Retirement Researcher) are cited directly by URL above. Treat figures quoted from "
    "secondary coverage (e.g. exact success-rate percentages) as approximate pending a check against "
    "the primary paper if they are to be quoted externally."
)
r.italic = True
r.font.size = Pt(9.5)
r.font.color.rgb = GREY

OUT.parent.mkdir(exist_ok=True)
doc.save(OUT)
print(f"Saved {OUT}")
