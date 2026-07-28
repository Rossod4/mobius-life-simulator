"""
Builds output/Mobius_Wealth_UK_Equity_Data_Extraction_Spec.docx - internship Week 5, Task 12
("Download UK equity/company-level data from Bloomberg"). A short, actionable companion to
build_equity_literature_review.py's Task 11 document and to extract_uk_equity_data.py, which does
the actual parsing once the export below exists.

Run: `python build_equity_data_spec.py`
"""
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "output" / "Mobius_Wealth_UK_Equity_Data_Extraction_Spec.docx"
CANDIDATES_CSV = Path(__file__).resolve().parent.parent / "data" / "equities" / "uk_shares_universe_candidates.csv"

NAVY = RGBColor(0x1A, 0x2B, 0x3C)
GREEN = RGBColor(0x1B, 0xAF, 0x7A)
GREY = RGBColor(0x5A, 0x5A, 0x5A)

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def title_block():
    p = doc.add_paragraph()
    r = p.add_run("Mobius Wealth")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    r = p.add_run("UK Equity Data Extraction Spec")
    r.font.size = Pt(20)
    r.font.color.rgb = GREEN

    p = doc.add_paragraph()
    r = p.add_run("Internship Week 5, Task 12")
    r.font.size = Pt(16)
    r.font.color.rgb = GREY
    r2 = p.add_run(
        "\nExactly what to pull from Bloomberg Terminal, in what format, so it drops straight into "
        "the existing Weeks 5-8 framework (src/equity_income.py) with no code changes - see "
        "src/extract_uk_equity_data.py, which parses the export this spec describes."
    )
    r2.font.size = Pt(11)
    doc.add_paragraph()


title_block()

doc.add_heading("1. Why this is needed", level=1)
doc.add_paragraph(
    "src/equity_income.py (individual-share/basket decumulation testing) is fully built and tested, "
    "but currently runs on 10 synthetic placeholder tickers (PH-NBF, PH-ARG, etc. - fictional "
    "companies, see generate_placeholder_equity_data.py) rather than real UK shares. Every result it "
    "currently produces proves the MECHANICS work; none of it is a real finding yet. This is the one "
    "remaining blocker before Weeks 5-8 produce anything presentable."
)

doc.add_heading("2. Candidate share universe", level=1)
doc.add_paragraph(
    "16 real, long-listed FTSE 100 shares across 6 sectors - enough spread for the basket search "
    "(task 14) to have genuine diversification to find, and deliberately including several companies "
    "that cut dividends in 2020 or earlier, so the dataset isn't survivorship-biased towards only the "
    "safest-looking names. Adjust freely (add/remove tickers) before pulling - this is a starting "
    "point, not a fixed requirement. Full list also saved as "
    "data/equities/uk_shares_universe_candidates.csv."
)
candidates = pd.read_csv(CANDIDATES_CSV)
table = doc.add_table(rows=1, cols=4)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
for i, col in enumerate(["Ticker", "Company", "Sector", "Note"]):
    hdr[i].text = col
for _, row in candidates.iterrows():
    cells = table.add_row().cells
    cells[0].text = row["Ticker"]
    cells[1].text = row["Company"]
    cells[2].text = row["Sector"]
    cells[3].text = row["Note"]
doc.add_paragraph()

doc.add_heading("3. What field to pull", level=1)
doc.add_paragraph(
    "Monthly TOTAL RETURN (dividends reinvested), GBP, NOT a bare price series - a price-only series "
    "understates income shares' real return and would defeat the point of testing 'equities for "
    "retirement income' in the first place. In the Bloomberg Excel Add-in this is typically a "
    "monthly-periodicity BDH() pull on a total-return field (e.g. TOT_RETURN_INDEX_GROSS_DVDS). "
    "Check with whoever ran the original asset-class-level pull (Bloomberg_Decumulation_Data_9_"
    "July_2026.xlsx) for the exact field/settings used there, so this new pull is methodologically "
    "consistent with the rest of the model rather than mixing conventions."
)

doc.add_heading("4. Date range", level=1)
doc.add_paragraph(
    "As far back as Bloomberg has data for each name - ideally to 1999-2000 to match the rest of "
    "this model's history. Several names will have less (e.g. any demergers/listings after 2000); "
    "that's fine, each share's own available history is used, the same way the main app already "
    "handles portfolios with different start dates."
)

doc.add_heading("5. Export format (must match exactly)", level=1)
doc.add_paragraph(
    "Lay the export out identically to the existing Bloomberg_Decumulation_Data file so the same "
    "parsing logic (extract_data.py's block parser, reused as-is in extract_uk_equity_data.py) "
    "works unchanged:"
)
for b in [
    "One 3-column block per ticker: column 1 = Date, column 2 = Monthly Return, column 3 = blank spacer.",
    "Row 1 of each block = a header identifying the ticker/company (e.g. 'ULVR LN Equity' or "
    "'Unilever plc') - extract_uk_equity_data.py auto-matches this back to the candidate list in "
    "Section 2, and flags anything it can't match for a manual fix.",
    "Row 2 = column labels (e.g. 'Date' / 'Monthly Return (Bloomberg direct)').",
    "Data from row 3 down.",
    "Save as .xlsx, sheet named 'Bloomberg Direct Returns' (or update SHEET_NAME at the top of "
    "extract_uk_equity_data.py if it's named differently).",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("6. Once the export exists", level=1)
for b in [
    "Open src/extract_uk_equity_data.py and point SRC at the saved file's path.",
    "Run: python extract_uk_equity_data.py",
    "This produces data/equities/uk_shares_returns_real.csv and "
    "data/equities/share_metadata_real.csv, in exactly the shape equity_income.py already expects.",
    "Point equity_income.py's EQUITY_RETURNS_CSV/SHARE_METADATA_CSV constants at the two _real "
    "files (or rename them over the placeholder ones), then re-run run_equity_income_analysis.py - "
    "rank_shares, find_best_baskets, evaluate_basket and share_correlation_matrix all work "
    "unchanged, now on real data.",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run(
    "Note: this document specifies the pull; it does not perform it. Bloomberg Terminal access is "
    "needed to actually run the extraction - see the companion literature review (Task 11) for the "
    "research rationale behind testing shares individually, in baskets, and across rebalancing "
    "conventions, which this data feeds into."
)
r.italic = True
r.font.size = Pt(9.5)
r.font.color.rgb = GREY

OUT.parent.mkdir(exist_ok=True)
doc.save(OUT)
print(f"Saved {OUT}")
