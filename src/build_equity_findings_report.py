"""
Builds output/Mobius_Wealth_UK_Equities_Weeks5-6_Findings.docx - internship Task 17 ("Write up and
present findings from Weeks 5-6" - the main deliverable). Styled to match
build_equity_literature_review.py / Mobius_Wealth_Methodology_and_Assumptions.docx.

Runs on REAL Bloomberg data (task 12 complete - see extract_uk_equity_data.py and
data/equities/uk_shares_returns_real.csv) rather than the earlier placeholder/synthetic series.

Run: `python build_equity_findings_report.py`
"""
from pathlib import Path

import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from engine import load_asset_returns, load_cpi, ClientProfile
from equity_income import (
    load_equity_returns, load_share_metadata, rank_shares, share_correlation_matrix,
    find_best_baskets, evaluate_basket,
)

OUT = Path(__file__).resolve().parent.parent / "output" / "Mobius_Wealth_UK_Equities_Weeks5-6_Findings.docx"
CHART_DIR = Path(__file__).resolve().parent.parent / "output" / "charts"
CHART_DIR.mkdir(parents=True, exist_ok=True)

SECTOR_COLORS = {
    "Consumer Staples": "#1BAF7A",
    "Healthcare": "#3C7DC4",
    "Financials": "#EDA100",
    "Utilities": "#6B6F76",
    "Mining": "#8B5E3C",
    "Energy": "#D03B3B",
}
RUIN_GOOD, RUIN_WARN, RUIN_BAD = "#1BAF7A", "#EDA100", "#D03B3B"


def ruin_color(p):
    return RUIN_GOOD if p < 0.10 else RUIN_WARN if p < 0.30 else RUIN_BAD

NAVY = RGBColor(0x1A, 0x2B, 0x3C)
GREEN = RGBColor(0x1B, 0xAF, 0x7A)
GREY = RGBColor(0x5A, 0x5A, 0x5A)
LIGHTGREEN_HEX = "E6F7EF"
LIGHTGREY_HEX = "F2F2F2"

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def callout(text, fill=LIGHTGREEN_HEX, border_color="1BAF7A"):
    """A shaded, bordered 1-cell 'table' used as a callout box - python-docx has no native box
    element, this is the standard workaround."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "12")
        el.set(qn("w:color"), border_color)
        borders.append(el)
    tcPr.append(borders)
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = NAVY
    doc.add_paragraph()


def title_block():
    p = doc.add_paragraph()
    r = p.add_run("Mobius Wealth")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    r = p.add_run("UK Equities for Retirement Income")
    r.font.size = Pt(20)
    r.font.color.rgb = GREEN

    p = doc.add_paragraph()
    r = p.add_run("Weeks 5-6 Findings")
    r.font.size = Pt(16)
    r.font.color.rgb = GREY
    r2 = p.add_run(
        "\nInternship Weeks 5-6, Tasks 13-16 - individual-share and basket-level decumulation "
        "testing (src/equity_income.py), extending the main Mobius Wealth asset-class simulator "
        "down to individual UK equities."
    )
    r2.font.size = Pt(11)
    p2 = doc.add_paragraph()
    r3 = p2.add_run("Hasini Yahampath")
    r3.italic = True
    r3.font.size = Pt(10)
    r3.font.color.rgb = GREY
    doc.add_paragraph()


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def body(text):
    doc.add_paragraph(text)


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def make_table(headers, rows, col_widths_cm=None, font_pt=None, shade_col=None, shade_fn=None):
    """shade_col: index of a column whose cells get a background fill from shade_fn(raw_value) -
    raw_value is the corresponding entry in `rows` before str() conversion, e.g. a raw 0-1
    probability, so the caller doesn't need to re-parse the formatted "12.34%" string."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                if font_pt:
                    run.font.size = Pt(font_pt)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            if font_pt:
                for p in cells[i].paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(font_pt)
            if shade_col is not None and i == shade_col and shade_fn is not None:
                shade_cell(cells[i], shade_fn(val).lstrip("#"))
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_chart(path, width_in=6.3, caption=None):
    doc.add_picture(str(path), width=Inches(width_in))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GREY
    doc.add_paragraph()


def start_landscape_section():
    """New page, switched to landscape - used for the wide 16x16 correlation matrix, which does not
    fit readably in portrait A4."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_w, new_h = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = new_w, new_h
    return section


def end_landscape_section():
    """Switches back to portrait for everything after the correlation matrix."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_w, new_h = section.page_height, section.page_width
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = new_w, new_h
    return section


# ============================================================================================
# Compute the real results (task 13, 14, 16) fresh, same as run_equity_income_analysis.py
# ============================================================================================
print("Computing results for the report (this re-runs the full Weeks 5-6 analysis)...")
asset_df = load_asset_returns()
cpi = load_cpi(asset_df)
equity_df = load_equity_returns()
meta = load_share_metadata()

profile = ClientProfile(starting_age=65, horizon_years=30, starting_pot=500_000.0, initial_annual_spend=20_000.0)

ranked = rank_shares(equity_df, cpi, profile)
ranked = ranked.merge(meta[["Ticker", "Company", "Sector"]], left_on="Share", right_on="Ticker")
ranked = ranked.sort_values("Probability of ruin")

sector_map = dict(zip(meta["Ticker"], meta["Sector"]))
SEARCH_N_SIMS = 300
top_baskets = find_best_baskets(equity_df, cpi, profile, basket_size=3, top_n=5,
                                 n_sims=SEARCH_N_SIMS, sector_map=sector_map, min_sectors=3)

best_basket = top_baskets.iloc[0]["Basket"]
best_tickers = best_basket.split(" + ")
weights = {t: 1.0 / len(best_tickers) for t in best_tickers}

rebal_rows = []
for label, mode in [("Constant-mix (rebalanced monthly)", "monthly"),
                     ("Annual rebalance", "annual"),
                     ("Buy-and-hold (never rebalanced)", "buy_and_hold")]:
    res, dd = evaluate_basket(f"Best basket ({mode})", weights, equity_df, cpi, profile, rebalance=mode)
    s = res.summary()
    rebal_rows.append((label, s["Probability of ruin"], s["Median legacy"], dd["maxdd"]))

data_start = equity_df.index.min().strftime("%b %Y")
data_end = equity_df.index.max().strftime("%b %Y")
n_shares = len(ranked)
best_share = ranked.iloc[0]
worst_share = ranked.iloc[-1]

# Characteristics: annualised return/vol (same convention as historical_stats_arithmetic elsewhere
# in the project - arithmetic monthly mean/std, compounded/scaled to annual) plus each share's own
# available data window, since not every share has the full Dec 1999-Jul 2026 history (Shell).
char_rows = []
for _, m in meta.iterrows():
    s = equity_df[m["Ticker"]].dropna()
    ann_return = (1 + s.mean()) ** 12 - 1
    ann_vol = s.std() * np.sqrt(12)
    char_rows.append({
        "Ticker": m["Ticker"], "Company": m["Company"], "Sector": m["Sector"],
        "Ann. return": ann_return, "Ann. vol": ann_vol,
        "Data period": f"{s.index.min().strftime('%b %Y')} - {s.index.max().strftime('%b %Y')}",
    })
characteristics = pd.DataFrame(char_rows).sort_values("Ann. return", ascending=False)

# NOTE: evaluate_basket() (used just above, for the Task 16 rebalancing comparison) mutates
# equity_df in place, adding a column per basket it registers - restrict to the original 16 share
# tickers here so the correlation matrix doesn't pick those basket columns up too.
corr = share_correlation_matrix(equity_df[list(meta["Ticker"])])

print("Results computed. Building charts...")


def _style_ax(ax):
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.tick_params(labelsize=8)


def _save(fig, name):
    path = CHART_DIR / name
    fig.tight_layout()
    fig.savefig(path, dpi=170)
    plt.close(fig)
    return path


# --- Chart 1: Task 13 individual share ranking ---
fig, ax = plt.subplots(figsize=(7, 5))
plot_df = ranked.sort_values("Probability of ruin", ascending=True)
colors = [SECTOR_COLORS.get(s, "#999999") for s in plot_df["Sector"]]
bars = ax.barh(plot_df["Share"], plot_df["Probability of ruin"] * 100, color=colors)
ax.invert_yaxis()
for bar, val in zip(bars, plot_df["Probability of ruin"]):
    ax.text(bar.get_width() + 1, bar.get_y() + bar.get_height() / 2, f"{val*100:.1f}%",
            va="center", fontsize=8)
ax.set_xlabel("Probability of ruin (%)")
ax.set_xlim(0, plot_df["Probability of ruin"].max() * 100 * 1.18)  # headroom for data labels
_style_ax(ax)
handles = [plt.Rectangle((0, 0), 1, 1, color=c) for c in SECTOR_COLORS.values()]
ax.legend(handles, SECTOR_COLORS.keys(), loc="upper center", bbox_to_anchor=(0.5, -0.12),
          fontsize=7.5, ncol=3, frameon=False)
chart1 = _save(fig, "task13_ranking.png")

# --- Chart 2: correlation heatmap ---
fig, ax = plt.subplots(figsize=(7, 6.5))
im = ax.imshow(corr.values, cmap="RdYlGn", vmin=-1, vmax=1)
ax.set_xticks(range(len(corr.columns)))
ax.set_xticklabels(corr.columns, rotation=90, fontsize=7)
ax.set_yticks(range(len(corr.index)))
ax.set_yticklabels(corr.index, fontsize=7)
for i in range(len(corr.index)):
    for j in range(len(corr.columns)):
        ax.text(j, i, f"{corr.values[i, j]:.1f}", ha="center", va="center", fontsize=5.5,
                color="black" if abs(corr.values[i, j]) < 0.7 else "white")
fig.colorbar(im, ax=ax, shrink=0.8, label="Correlation")
chart2 = _save(fig, "correlation_heatmap.png")

# --- Chart 3: characteristics scatter (return vs vol by sector) ---
fig, ax = plt.subplots(figsize=(7, 5.5))
for sector, grp in characteristics.groupby("Sector"):
    ax.scatter(grp["Ann. vol"] * 100, grp["Ann. return"] * 100, s=90,
               color=SECTOR_COLORS.get(sector, "#999999"), label=sector, edgecolor="white", linewidth=0.6)
    for _, row in grp.iterrows():
        ax.annotate(row["Ticker"], (row["Ann. vol"] * 100, row["Ann. return"] * 100),
                    fontsize=7.5, xytext=(4, 4), textcoords="offset points")
ax.set_xlabel("Annualised volatility (%)")
ax.set_ylabel("Annualised return (%)")
ax.legend(fontsize=7.5, frameon=False, loc="upper left")
_style_ax(ax)
chart3 = _save(fig, "characteristics_scatter.png")

# --- Chart 4: Task 14 top baskets ---
fig, ax = plt.subplots(figsize=(7, 3.5))
plot_b = top_baskets.sort_values("Probability of ruin", ascending=True)
bars = ax.barh(plot_b["Basket"], plot_b["Probability of ruin"] * 100, color="#1BAF7A")
ax.invert_yaxis()
for bar, val in zip(bars, plot_b["Probability of ruin"]):
    ax.text(bar.get_width() + 0.05, bar.get_y() + bar.get_height() / 2, f"{val*100:.2f}%",
            va="center", fontsize=8)
ax.set_xlabel("Probability of ruin (%)")
_style_ax(ax)
chart4 = _save(fig, "task14_baskets.png")

# --- Chart 5: Task 16 rebalancing comparison ---
fig, ax = plt.subplots(figsize=(6, 3.5))
labels5 = [l.split(" (")[0] for l, *_ in rebal_rows]
vals5 = [p * 100 for _, p, _, _ in rebal_rows]
bars = ax.bar(labels5, vals5, color=["#1BAF7A", "#3C7DC4", "#EDA100"])
for bar, val in zip(bars, vals5):
    ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.005, f"{val:.2f}%",
            ha="center", fontsize=8)
ax.set_ylabel("Probability of ruin (%)")
_style_ax(ax)
chart5 = _save(fig, "task16_rebalancing.png")

print("Charts built. Building document...")

# ============================================================================================
# Document
# ============================================================================================
title_block()

callout(
    "Data provenance: this report runs on REAL Bloomberg data (Task 12, now complete) - monthly "
    f"total return (dividends reinvested, GBP) for {n_shares} FTSE 100 shares, {data_start} to "
    f"{data_end} (Shell's shorter ~2005-{data_end} history reflects its own 2022 dual-listing "
    "unification under this ticker, not a data error). This supersedes the earlier placeholder/"
    "synthetic-data version of this analysis - every number below is a real, presentable finding, "
    "not a mechanics-only demonstration."
)

h1("Executive summary")
body(
    f"Every one of the {n_shares} candidate FTSE 100 shares was tested alone, then systematically "
    "searched in every possible 3-share, sector-diversified combination, against the same "
    "objective the rest of this project uses: the probability a £500,000 pot drawing £20,000/year "
    "runs out within 30 years. The best single share on its own "
    f"({best_share['Share']}, {best_share['Company']}) already achieves a "
    f"{best_share['Probability of ruin']*100:.1f}% probability of ruin - comparable to Mobius "
    "Better, the main simulator's own diversified multi-asset benchmark (~3.6%). The best "
    f"3-share baskets go further still, reaching a probability of ruin of "
    f"{top_baskets.iloc[0]['Probability of ruin']*100:.1f}% - a genuinely striking result, though "
    "one that needs the caveats in the Discussion section below before it is presented any further."
)

h1("Background & objective")
body(
    "This work extends the main Mobius Wealth decumulation simulator - which tests portfolios at "
    "the asset-class level (bonds, global equities, hedge strategies, etc.) - down to individual UK "
    "shares. The rationale for testing equities for retirement income at all, and for using "
    "probability of ruin (rather than average return) as the objective function, is set out in the "
    "companion document, 'Equities for Retirement Income: Literature Review' (Task 11) - "
    "sequence-of-returns risk in particular is why probability of ruin, not average return, is the "
    "headline metric throughout."
)

h1("Methodology")
h2("Objective function")
body(
    "Same convention as the rest of the project: a Monte Carlo, historical-bootstrap simulation of "
    "a £500,000 pot drawing £20,000/year (4% initial withdrawal rate) for a 65-year-old over a "
    "30-year horizon (src/engine.py's run_simulation), with 'ruin' defined as the pot reaching £0 "
    "before the 30 years are up."
)
h2("Task 13 - individual share testing")
body(
    "Each of the 16 candidate shares was tested entirely alone (100% weight, no diversification) "
    "against the objective function above, using rank_shares() in src/equity_income.py."
)
h2("Task 14 - basket search")
body(
    "A systematic search evaluates every possible 3-share equal-weight combination - not a "
    "hand-picked pair - restricted to baskets spanning 3 distinct GICS-style sectors, to guarantee "
    "genuine sector diversification rather than diversification that only looks real over one "
    "historical sample (two shares can show low historical correlation over a ~25-year window by "
    "chance without being economically diversified at all). The candidate winner is screened at a "
    f"reduced simulation count (n_sims={SEARCH_N_SIMS}) for speed across the full combinatorial "
    "search, then re-verified at full simulation count in Task 16 below."
)
h2("Task 16 - rebalancing approaches")
body(
    "The winning basket from Task 14 is re-tested across three rebalancing conventions: "
    "constant-mix (rebalanced back to equal weight every month), annual rebalance, and "
    "buy-and-hold (never rebalanced, weights drift with performance)."
)

h1("Results")
h2("Task 13 - individual share ranking")
headers = ["Share", "Company", "Sector", "Prob. of ruin", "Median legacy", "Max DD", "Avg DD", "CVaR 95 Mthly"]
rows = [
    (r["Share"], r["Company"], r["Sector"], f"{r['Probability of ruin']*100:.2f}%",
     f"£{r['Median legacy']:,.0f}", f"{r['Max DD']*100:.2f}%", f"{r['Average DD']*100:.2f}%",
     f"{r['CVaR 95 Mthly']*100:.2f}%")
    for _, r in ranked.iterrows()
]
make_table(headers, rows, shade_col=3, shade_fn=lambda v: ruin_color(float(v.rstrip("%")) / 100))
add_chart(chart1, caption="Probability of ruin by share, coloured by sector (lower is safer).")
p = doc.add_paragraph()
r = p.add_run(
    f"Note on the largest median legacy figures (BATS £{ranked[ranked['Share']=='BATS']['Median legacy'].iloc[0]:,.0f}, "
    "RIO similarly large): these are genuine outputs of the bootstrap, not errors. BAT's ~18.7% "
    "annualised return over this 25-year window, compounded with no rebalancing/diversification "
    "for a full 30-year simulated horizon from a single concentrated position, mechanically "
    "produces very large figures. They should NOT be read as an expected or likely outcome for a "
    "real investor - a single-share position carries far more variance than these median-outcome "
    "figures alone convey (see the Discussion section's concentration-risk caveat), and reflects "
    "this one company's specific historical run, not a repeatable property of holding one share."
)
r.italic = True
r.font.size = Pt(9.5)
r.font.color.rgb = GREY
body(
    f"{best_share['Share']} ({best_share['Company']}) is the standout single name, comparable to "
    f"Mobius Better's diversified multi-asset benchmark on this metric alone. "
    f"{worst_share['Share']} ({worst_share['Company']}) is the weakest "
    f"({worst_share['Probability of ruin']*100:.1f}% probability of ruin), consistent with its "
    "2020 dividend cut and higher realised volatility through the Global Financial Crisis. The "
    "spread between best and worst single shares is wide - underlining that concentrated, "
    "single-name exposure carries real dispersion of outcomes even before considering "
    "diversification, which motivates the basket search below."
)

h2("Task 14 - top 5 baskets found")
headers2 = ["Basket", "Prob. of ruin", "Median legacy", "Max DD", "Avg DD", "Sectors"]
rows2 = [
    (r["Basket"], f"{r['Probability of ruin']*100:.2f}%", f"£{r['Median legacy']:,.0f}",
     f"{r['Max DD']*100:.2f}%", f"{r['Average DD']*100:.2f}%", r["Sectors"])
    for _, r in top_baskets.iterrows()
]
make_table(headers2, rows2)
add_chart(chart4, width_in=6.3, caption="Probability of ruin for the top 5 sector-diversified 3-share baskets found.")
body(
    f"The winning basket ({best_basket}) combines defensive Consumer Staples/Healthcare/Utilities "
    "exposure with Energy - all five shortlisted baskets follow a similar pattern of pairing "
    "defensive sectors with one more cyclical name. Combining into a 3-way sector-diversified "
    "basket improves meaningfully on any single share alone, and - notably - several of these "
    "baskets edge out Mobius Better's own diversified multi-asset benchmark (~3.6% probability of "
    "ruin) on this specific historical window. See the Discussion section for why this should not "
    "yet be read as 'a 3-share basket beats a properly diversified portfolio' as a general claim."
)

h2("Task 16 - rebalancing comparison for the winning basket")
headers3 = ["Approach", "Prob. of ruin", "Median legacy", "Max DD"]
rows3 = [(label, f"{p*100:.2f}%", f"£{ml:,.0f}", f"{dd*100:.2f}%") for label, p, ml, dd in rebal_rows]
make_table(headers3, rows3)
add_chart(chart5, width_in=5.2, caption="Rebalancing convention vs probability of ruin, winning basket.")
body(
    "Constant-mix and annual rebalancing perform similarly here and both edge out buy-and-hold on "
    "probability of ruin, consistent with rebalancing's known benefit of systematically trimming "
    "winners and topping up laggards rather than letting one position's drift dominate the "
    "basket's risk profile over time."
)

h1("Share characteristics and correlation")
h2("Characteristics")
body(
    "Annualised return and volatility for each share individually (arithmetic monthly mean/std, "
    "same convention used for annualised performance figures throughout this project), alongside "
    "each share's own available data window - not every share has the full Dec 1999-Jul 2026 "
    "history."
)
char_headers = ["Ticker", "Company", "Sector", "Ann. return", "Ann. vol", "Data period"]
char_table_rows = [
    (r["Ticker"], r["Company"], r["Sector"], f"{r['Ann. return']*100:.2f}% pa",
     f"{r['Ann. vol']*100:.2f}% pa", r["Data period"])
    for _, r in characteristics.iterrows()
]
make_table(char_headers, char_table_rows)
add_chart(chart3, width_in=6.3, caption="Annualised return vs volatility by share, coloured by sector.")
body(
    "Return and volatility here are UNCONDITIONAL characteristics of each share's own return "
    "series - they do not by themselves determine the probability-of-ruin ranking in the Task 13 "
    "table above, which also depends on the sequence and timing of returns relative to the "
    "withdrawal schedule (sequence-of-returns risk, per the companion literature review), not just "
    "the average return/volatility level."
)

h2("Pairwise correlation")
body(
    "Full pairwise correlation of monthly returns across all 16 shares - used alongside the "
    "sector-spread basis for Task 14's basket search (see Methodology) to understand WHY a "
    "sector-diverse basket behaves the way it does, rather than as the primary diversification "
    "test itself. The financials cluster (HSBA/LLOY/LGEN/AV.) and the miners (RIO/AAL) and oil "
    "majors (SHEL/BP.) show the highest same-sector correlation, as expected; cross-sector pairs "
    "are consistently lower, supporting the sector-based diversification approach used in Task 14."
)
add_chart(chart2, width_in=6.3, caption="Pairwise correlation heatmap - green/red = higher positive/negative correlation.")
body("Exact figures, for reference:")
start_landscape_section()
corr_headers = [""] + list(corr.columns)
corr_rows = [[idx] + [f"{v:.2f}" for v in row] for idx, row in zip(corr.index, corr.values)]
make_table(corr_headers, corr_rows, font_pt=7)
end_landscape_section()

h1("Discussion")
body(
    "The headline numbers above are real and are a genuine result of this specific historical "
    "backtest - not a placeholder or a mechanics demonstration. Several things are worth keeping "
    "in view before presenting them any further, though:"
)
bullet(
    "Small universe, small basket: 16 candidate shares and 3-share baskets is a narrow slice of "
    "the market. A ~0% probability of ruin on a 3-name basket reflects how well THESE specific "
    "names happened to perform over THIS specific ~25-year UK-only window - it is not the same "
    "statistical footing as Mobius Better, which spreads risk across many more holdings, asset "
    "classes, and geographies."
)
bullet(
    "Single-name/concentration risk is real even when the backtest looks strong: a 3-share basket "
    "still carries meaningful company-specific risk (an accounting scandal, a takeover, a "
    "regulatory shock) that a professionally diversified multi-asset portfolio structurally "
    "diversifies away. Past performance not showing this risk crystallising is not the same as the "
    "risk not existing."
)
bullet(
    "Survivorship and hindsight: the search finds whichever combination worked best over the "
    "sample with the benefit of hindsight. An investor choosing a basket in real time in 1999 could "
    "not have known which shares would perform this way over the following 25 years."
)
bullet(
    "Shell's shorter data history (~2005 onward under this ticker) means its contribution to any "
    "basket including it is tested over a shorter, more recent window than the other names."
)
body(
    "None of this means the result is wrong or uninteresting - it is a legitimate finding that "
    "sector-diversified UK equity baskets can perform strongly against this objective over this "
    "period, and a reasonable basis for further work. It does mean the result should be framed as "
    "'a promising finding worth testing further' (Task 18) rather than 'a 3-share basket is a "
    "substitute for a properly diversified multi-asset portfolio' as a general conclusion."
)

h1("Next steps")
bullet("Task 12 (real Bloomberg data) is now complete - this document supersedes the earlier placeholder-data version.")
bullet(
    "Task 18: incorporate feedback and test alternative approaches - e.g. larger baskets (4-5 "
    "shares), a wider candidate universe beyond the original 16, and testing sensitivity to the "
    "exact historical window used."
)
bullet("Task 19 (Week 8): finalise the report incorporating any feedback from Task 18.")
bullet("Task 20 (optional, if time permits): extend the same framework to global developed equity markets (e.g. US, Europe).")

h1("Appendix")
body(
    "Companion documents: 'Equities for Retirement Income: Literature Review' (Task 11), 'UK "
    "Equity Data Extraction Spec' (Task 12). Source: src/equity_income.py (framework), "
    "src/run_equity_income_analysis.py (this analysis), src/extract_uk_equity_data.py (Bloomberg "
    "export parser - handles the mixed percentage-point/decimal scale found across this export's "
    "columns automatically, see its _fix_scale function)."
)

p = doc.add_paragraph()
r = p.add_run(
    "Note: probability-of-ruin and median-legacy figures are Monte Carlo estimates from a "
    "historical bootstrap over the available data window for each share/basket, not a guarantee of "
    "future performance. This is an internal working document for the internship project, not a "
    "client-facing deliverable."
)
r.italic = True
r.font.size = Pt(9.5)
r.font.color.rgb = GREY

OUT.parent.mkdir(exist_ok=True)
try:
    doc.save(OUT)
    print(f"Saved {OUT}")
except PermissionError:
    fallback = OUT.with_stem(OUT.stem + "_v2")
    doc.save(fallback)
    print(f"{OUT} is locked (likely open in Word) - saved to {fallback} instead. "
          f"Close the original and re-run to overwrite it directly next time.")
